from rich.progress import Progress
from ocr import process_image
from io import BytesIO
from PIL import Image
import pymupdf
import logging
import transformers
import warnings
from docx import Document
from docx.shared import Pt

warnings.filterwarnings("ignore", category=UserWarning, module="transformers")
transformers.logging.set_verbosity_error()
logging.getLogger("transformers").setLevel(logging.ERROR)

doc = pymupdf.open("moaser.pdf")
word_doc = Document()
word_doc.add_heading("OCR Output", 0)


with Progress() as progress:
    task = progress.add_task("", total=len(doc))
    for page_index in range(len(doc)):
        page = doc[page_index]
        image_list = page.get_images()
        for image_index, img in enumerate(image_list, start=1):
            xref = img[0]
            pix = pymupdf.Pixmap(doc, xref)
            if pix.n - pix.alpha > 3:
                pix = pymupdf.Pixmap(pymupdf.csRGB, pix)
            img_bytes = pix.tobytes("png")
            pil_img = Image.open(BytesIO(img_bytes))

            cleaned, markdown, raw, img_out, crops = process_image(
                pil_img, task="📋 Markdown", custom_prompt=""
            )

            word_doc.add_heading(f"Page {page_index + 1}", level=1)
            word_doc.add_paragraph(cleaned)
            pix = None
        progress.update(task, advance=1)

word_doc.save("output.docx")
